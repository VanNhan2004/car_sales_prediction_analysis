# report_docx.py
# =====================================================
# TẠO BÁO CÁO WORD (.DOCX) DỰ ĐOÁN NHU CẦU MUA XE
# =====================================================

from io import BytesIO
from datetime import datetime
from docx import Document
from docx.shared import Pt


def export_prediction_docx(
    input_data: dict,
    prediction: float,
    shap_df,
    llm_analysis: str
):
    """
    Xuất báo cáo dự đoán ra file Word (.docx)
    Trả về buffer để Streamlit download
    """

    doc = Document()

    # =========================
    # STYLE
    # =========================
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    # =========================
    # TIÊU ĐỀ
    # =========================
    title = doc.add_heading("BÁO CÁO DỰ ĐOÁN NHU CẦU MUA XE", level=1)
    title.alignment = 1  # center

    doc.add_paragraph(
        f"Thời gian tạo báo cáo: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}"
    )

    doc.add_paragraph("")

    # =========================
    # 1. THÔNG TIN CẤU HÌNH XE
    # =========================
    doc.add_heading("1. Thông tin cấu hình xe", level=2)

    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Thuộc tính"
    hdr_cells[1].text = "Giá trị"

    for k, v in input_data.items():
        row_cells = table.add_row().cells
        row_cells[0].text = str(k)
        row_cells[1].text = str(v)

    # =========================
    # 2. KẾT QUẢ DỰ ĐOÁN
    # =========================
    doc.add_paragraph("")
    doc.add_heading("2. Kết quả dự đoán", level=2)

    doc.add_paragraph(
        f"Tổng số người mua dự đoán: {int(prediction):,}"
    )

    # =========================
    # 3. CÁC YẾU TỐ ẢNH HƯỞNG CHÍNH
    # =========================
    doc.add_paragraph("")
    doc.add_heading("3. Các yếu tố ảnh hưởng chính", level=2)

    shap_table = doc.add_table(rows=1, cols=3)
    shap_table.style = "Table Grid"
    hdr = shap_table.rows[0].cells
    hdr[0].text = "Yếu tố"
    hdr[1].text = "Giá trị"
    hdr[2].text = "Mức ảnh hưởng (SHAP)"

    for _, row in shap_df.iterrows():
        r = shap_table.add_row().cells
        r[0].text = str(row["feature"])
        r[1].text = str(row["value"])
        r[2].text = f"{row['shap_value']:,.2f}"

    # =========================
    # 4. KHUYẾN NGHỊ KINH DOANH
    # =========================
    doc.add_paragraph("")
    doc.add_heading("4. Khuyến nghị kinh doanh", level=2)

    for line in llm_analysis.split("\n"):
        if line.strip():
            doc.add_paragraph(line)

    # =========================
    # EXPORT BUFFER
    # =========================
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer
